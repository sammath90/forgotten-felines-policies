#!/usr/bin/env python3
"""Rebuild .docx versions of the Forgotten Felines policy suite from the
current Markdown sources.

Supports a focused subset of GitHub-flavoured Markdown sufficient for the
Forgotten Felines policies: headings, paragraphs, bullet / numbered lists,
GFM pipe tables, fenced code blocks, blockquotes, horizontal rules, and
inline **bold**, *italic*, `code`, [link text](url).

Run:
    python3 scripts/md_to_docx.py
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError:  # pragma: no cover
    print("python-docx is required: pip install python-docx", file=sys.stderr)
    sys.exit(1)


REPO = Path(__file__).resolve().parent.parent
PUBLISHED = REPO / "docs/forgotten-felines/published-policies"

JOBS: list[tuple[str, str]] = [
    # Animal-care v2.0 policies
    ("docs/forgotten-felines/policies/animal-care/adoption-policy.md",
     "adoption-policy.docx"),
    ("docs/forgotten-felines/policies/animal-care/cat-fosterer-cleaning-policy.md",
     "cat-fosterer-cleaning-policy.docx"),
    ("docs/forgotten-felines/policies/animal-care/cat-intake-policy.md",
     "cat-intake-policy.docx"),
    ("docs/forgotten-felines/policies/animal-care/cat-isolation-arrival-policy.md",
     "cat-isolation-arrival-policy.docx"),
    ("docs/forgotten-felines/policies/animal-care/cat-relinquishment-policy.md",
     "cat-relinquishment-policy.docx"),
    ("docs/forgotten-felines/policies/animal-care/cat-rescue-euthanasia-policy.md",
     "cat-rescue-euthanasia-policy.docx"),
    ("docs/forgotten-felines/policies/animal-care/cat-transportation-policy.md",
     "cat-transportation-policy.docx"),
    ("docs/forgotten-felines/policies/animal-care/death-of-a-cat-policy.md",
     "death-of-a-cat-policy.docx"),
    ("docs/forgotten-felines/policies/animal-care/escape-of-a-cat-policy.md",
     "escape-of-a-cat-policy.docx"),
    ("docs/forgotten-felines/policies/animal-care/fostering-policy.md",
     "fostering-policy.docx"),
    ("docs/forgotten-felines/policies/animal-care/forgotten-felines-emergency-plan.md",
     "forgotten-felines-emergency-plan.docx"),
    # Visits and volunteers v2.0
    ("docs/forgotten-felines/policies/visits-and-volunteers/forgotten-felines-adopter-visit-policy.md",
     "forgotten-felines-adopter-visit-policy.docx"),
    ("docs/forgotten-felines/policies/visits-and-volunteers/forgotten-felines-fosterer-visit-policy.md",
     "forgotten-felines-fosterer-visit-policy.docx"),
    ("docs/forgotten-felines/policies/visits-and-volunteers/volunteer-policy.md",
     "volunteer-policy.docx"),
    ("docs/forgotten-felines/policies/visits-and-volunteers/social-media-policy.md",
     "social-media-policy.docx"),
    # Registers (also distributed as .docx for assessors)
    ("docs/forgotten-felines/registers/foster-home-assessment-checklist.md",
     "foster-home-assessment-checklist.docx"),
    ("docs/forgotten-felines/registers/foster-home-register.md",
     "foster-home-register.docx"),
    ("docs/forgotten-felines/registers/named-veterinary-practices-register.md",
     "named-veterinary-practices-register.docx"),
    ("docs/forgotten-felines/registers/per-cat-record-specification.md",
     "per-cat-record-specification.docx"),
    ("docs/forgotten-felines/registers/roles-and-responsibilities-register.md",
     "roles-and-responsibilities-register.docx"),
    # Templates
    ("docs/forgotten-felines/templates/fosterer-training-programme.md",
     "fosterer-training-programme.docx"),
    ("docs/forgotten-felines/templates/vet-arrangement-letter-template.md",
     "vet-arrangement-letter-template.docx"),
    # Governance
    ("docs/forgotten-felines/policies/governance/forgotten-felines-bullying-policy.md",
     "forgotten-felines-bullying-policy.docx"),
    ("docs/forgotten-felines/policies/governance/forgotten-felines-harassment-policy.md",
     "forgotten-felines-harassment-policy.docx"),
    ("docs/forgotten-felines/policies/governance/forgotten-felines-financial-policy.md",
     "forgotten-felines-financial-policy.docx"),
    ("docs/forgotten-felines/policies/governance/forgotten-felines-trustee-code-of-conduct-and-governance-policy.md",
     "forgotten-felines-trustee-code-of-conduct-and-governance-policy.docx"),
    # Data protection
    ("docs/forgotten-felines/policies/data-protection/forgotten-felines-data-retention-policy.md",
     "forgotten-felines-data-retention-policy.docx"),
    ("docs/forgotten-felines/policies/data-protection/forgotten-felines-gdpr-trustee-policy.md",
     "forgotten-felines-gdpr-trustee-policy.docx"),
    ("docs/forgotten-felines/policies/data-protection/forgotten-felines-gdpr-volunteer-policy.md",
     "forgotten-felines-gdpr-volunteer-policy.docx"),
    ("docs/forgotten-felines/policies/data-protection/forgotten-felines-privacy-notice.md",
     "forgotten-felines-privacy-notice.docx"),
    ("docs/forgotten-felines/policies/data-protection/forgotten-felines-record-of-processing-activities.md",
     "forgotten-felines-record-of-processing-activities.docx"),
    ("docs/forgotten-felines/policies/data-protection/forgotten-felines-data-breach-response-procedure.md",
     "forgotten-felines-data-breach-response-procedure.docx"),
    # Phase 1 deliverables — checklist distributed alongside the suite
    ("docs/phase-1-deliverables/phase-1-finalisation-checklist.md",
     "phase-1-finalisation-checklist.docx"),
]


# --- Inline formatting ---------------------------------------------------

INLINE_CODE = re.compile(r"`([^`]+)`")
BOLD = re.compile(r"\*\*([^*]+)\*\*")
ITALIC = re.compile(r"(?<!\*)\*([^*\n]+)\*(?!\*)")
LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def add_inline(paragraph, text: str) -> None:
    """Add text to a paragraph, applying inline Markdown formatting.

    Supported tokens, earliest match wins: `code`, **bold**, *italic*,
    [link text](url). Unknown characters fall through as plain text.
    """
    i = 0
    while i < len(text):
        candidates = []
        for pat, kind in (
            (INLINE_CODE, "code"),
            (BOLD, "bold"),
            (LINK, "link"),
            (ITALIC, "italic"),
        ):
            m = pat.search(text, i)
            if m:
                candidates.append((m.start(), m, kind))
        if not candidates:
            paragraph.add_run(text[i:])
            return
        candidates.sort(key=lambda x: x[0])
        start, m, kind = candidates[0]
        if start > i:
            paragraph.add_run(text[i:start])
        if kind == "code":
            run = paragraph.add_run(m.group(1))
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        elif kind == "bold":
            run = paragraph.add_run(m.group(1))
            run.bold = True
        elif kind == "italic":
            run = paragraph.add_run(m.group(1))
            run.italic = True
        elif kind == "link":
            run = paragraph.add_run(m.group(1))
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            run.underline = True
        i = m.end()


# --- Block parsing -------------------------------------------------------

HR_LINES = {"---", "***", "___"}


def is_block_start(line: str) -> bool:
    s = line.rstrip()
    if not s:
        return True
    if s.startswith("#"):
        return True
    if s in HR_LINES:
        return True
    if s.startswith("```"):
        return True
    if s.startswith(">"):
        return True
    if re.match(r"^[-*+]\s+", s):
        return True
    if re.match(r"^\d+\.\s+", s):
        return True
    if "|" in s and s.count("|") >= 2:
        return True
    return False


def split_table_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [cell.strip() for cell in s.split("|")]


def parse_blocks(lines: list[str]):
    """Yield block tuples for the rendered document."""
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.rstrip()

        if not stripped:
            i += 1
            continue

        # Fenced code block
        if stripped.startswith("```"):
            i += 1
            code_lines: list[str] = []
            while i < n and not lines[i].rstrip().startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            i += 1  # skip closing fence
            yield ("code", "\n".join(code_lines))
            continue

        # Horizontal rule
        if stripped in HR_LINES:
            yield ("hr",)
            i += 1
            continue

        # Heading
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            yield ("h", level, m.group(2))
            i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            quote_lines: list[str] = []
            while i < n and lines[i].rstrip().startswith(">"):
                raw = lines[i].rstrip()
                body = raw[1:].lstrip()
                quote_lines.append(body)
                i += 1
            yield ("quote", " ".join(filter(None, quote_lines)))
            continue

        # GFM table (requires a header row and a separator row)
        if (
            "|" in stripped
            and i + 1 < n
            and re.match(
                r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$",
                lines[i + 1].strip(),
            )
        ):
            rows: list[list[str]] = [split_table_row(stripped)]
            i += 2  # skip header + separator
            while i < n and "|" in lines[i] and lines[i].strip():
                rows.append(split_table_row(lines[i].rstrip()))
                i += 1
            yield ("table", rows)
            continue

        # Unordered list
        if re.match(r"^[-*+]\s+", stripped):
            items: list[str] = []
            while i < n and re.match(r"^[-*+]\s+", (lines[i].rstrip() or "")):
                items.append(re.sub(r"^[-*+]\s+", "", lines[i].rstrip()))
                i += 1
            yield ("ul", items)
            continue

        # Ordered list
        if re.match(r"^\d+\.\s+", stripped):
            ol_items: list[str] = []
            while i < n and re.match(r"^\d+\.\s+", (lines[i].rstrip() or "")):
                ol_items.append(re.sub(r"^\d+\.\s+", "", lines[i].rstrip()))
                i += 1
            yield ("ol", ol_items)
            continue

        # Paragraph: collect until blank line or new block start
        para_lines = [stripped]
        i += 1
        while i < n and lines[i].strip() and not is_block_start(lines[i]):
            para_lines.append(lines[i].rstrip())
            i += 1
        yield ("p", " ".join(para_lines))


# --- Rendering -----------------------------------------------------------

def render_docx(blocks, out_path: Path) -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for block in blocks:
        kind = block[0]

        if kind == "h":
            level, text = block[1], block[2]
            heading = doc.add_heading("", level=min(level, 4))
            add_inline(heading, text)

        elif kind == "p":
            p = doc.add_paragraph()
            add_inline(p, block[1])

        elif kind == "ul":
            for item in block[1]:
                p = doc.add_paragraph(style="List Bullet")
                add_inline(p, item)

        elif kind == "ol":
            for item in block[1]:
                p = doc.add_paragraph(style="List Number")
                add_inline(p, item)

        elif kind == "hr":
            p = doc.add_paragraph()
            run = p.add_run("_" * 40)
            run.italic = True

        elif kind == "quote":
            try:
                p = doc.add_paragraph(style="Intense Quote")
            except KeyError:
                p = doc.add_paragraph()
                for r in p.runs:
                    r.italic = True
            add_inline(p, block[1])

        elif kind == "code":
            p = doc.add_paragraph()
            run = p.add_run(block[1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)

        elif kind == "table":
            rows = block[1]
            if not rows:
                continue
            ncols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=ncols)
            for style_name in (
                "Light Grid Accent 1",
                "Light Grid",
                "Table Grid",
            ):
                try:
                    table.style = style_name
                    break
                except KeyError:
                    continue
            for r_idx, row in enumerate(rows):
                for c_idx in range(ncols):
                    cell_text = row[c_idx] if c_idx < len(row) else ""
                    cell = table.rows[r_idx].cells[c_idx]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    add_inline(p, cell_text)
                    if r_idx == 0:
                        for r in p.runs:
                            r.bold = True

    doc.save(str(out_path))


# --- Entry point ---------------------------------------------------------

def main() -> int:
    if not PUBLISHED.exists():
        PUBLISHED.mkdir(parents=True)
    rebuilt = 0
    skipped = 0
    for src_rel, out_name in JOBS:
        src = REPO / src_rel
        if not src.exists():
            print(f"SKIP missing: {src_rel}")
            skipped += 1
            continue
        lines = src.read_text(encoding="utf-8").splitlines()
        blocks = list(parse_blocks(lines))
        out = PUBLISHED / out_name
        render_docx(blocks, out)
        print(f"OK {out_name}")
        rebuilt += 1
    print(f"\n{rebuilt} rebuilt, {skipped} skipped")
    return 0


if __name__ == "__main__":
    sys.exit(main())
